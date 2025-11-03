"""
Google Drive Connector and File Parsers for Reasoner AI Platform.

Handles:
- Google Drive authentication and file listing
- File downloading and parsing (PDF, DOCX, XLSX, CSV)
- Data extraction for formula inputs
- Periodic sync operations
"""
import io
import json
from datetime import datetime
from typing import Dict, Any, List, Optional, Tuple
from pathlib import Path

from google.oauth2.credentials import Credentials
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from loguru import logger

# File parsers
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import openpyxl


class GoogleDriveConnector:
    """
    Connect to Google Drive and sync files for data ingestion.
    """
    
    def __init__(
        self,
        credentials_path: str,
        folder_id: Optional[str] = None,
        scopes: List[str] = None
    ):
        """
        Initialize Google Drive connector.
        
        Args:
            credentials_path: Path to service account JSON or OAuth credentials
            folder_id: Optional specific folder to monitor
            scopes: API scopes (default: read-only Drive access)
        """
        self.credentials_path = credentials_path
        self.folder_id = folder_id
        self.scopes = scopes or [
            'https://www.googleapis.com/auth/drive.readonly'
        ]
        
        self.service = None
        self._authenticate()
    
    def _authenticate(self):
        """Authenticate with Google Drive API."""
        try:
            # Try service account authentication
            credentials = service_account.Credentials.from_service_account_file(
                self.credentials_path,
                scopes=self.scopes
            )
            
            self.service = build('drive', 'v3', credentials=credentials)
            logger.info("Google Drive authenticated successfully")
            
        except Exception as e:
            logger.error(f"Google Drive authentication failed: {e}")
            raise
    
    def list_files(
        self,
        file_types: Optional[List[str]] = None,
        modified_after: Optional[datetime] = None,
        limit: int = 100
    ) -> List[Dict[str, Any]]:
        """
        List files in Google Drive.
        
        Args:
            file_types: Filter by mime types (e.g., ['application/pdf'])
            modified_after: Only return files modified after this date
            limit: Maximum number of files to return
            
        Returns:
            List of file metadata dicts
        """
        try:
            query_parts = []
            
            # Folder filter
            if self.folder_id:
                query_parts.append(f"'{self.folder_id}' in parents")
            
            # File type filter
            if file_types:
                mime_queries = [f"mimeType='{mt}'" for mt in file_types]
                query_parts.append(f"({' or '.join(mime_queries)})")
            
            # Modified date filter
            if modified_after:
                date_str = modified_after.isoformat()
                query_parts.append(f"modifiedTime > '{date_str}'")
            
            # Not trashed
            query_parts.append("trashed=false")
            
            query = " and ".join(query_parts)
            
            results = self.service.files().list(
                q=query,
                pageSize=limit,
                fields="files(id, name, mimeType, modifiedTime, size, webViewLink)",
                orderBy="modifiedTime desc"
            ).execute()
            
            files = results.get('files', [])
            logger.info(f"Found {len(files)} files in Google Drive")
            
            return files
            
        except Exception as e:
            logger.error(f"Failed to list Google Drive files: {e}")
            return []
    
    def download_file(self, file_id: str) -> bytes:
        """
        Download file content from Google Drive.
        
        Args:
            file_id: Google Drive file ID
            
        Returns:
            File content as bytes
        """
        try:
            request = self.service.files().get_media(fileId=file_id)
            
            file_content = io.BytesIO()
            downloader = MediaIoBaseDownload(file_content, request)
            
            done = False
            while not done:
                status, done = downloader.next_chunk()
                if status:
                    logger.debug(f"Download progress: {int(status.progress() * 100)}%")
            
            file_content.seek(0)
            return file_content.read()
            
        except Exception as e:
            logger.error(f"Failed to download file {file_id}: {e}")
            raise
    
    def sync_folder(
        self,
        local_cache_dir: str,
        file_types: Optional[List[str]] = None
    ) -> List[Dict[str, Any]]:
        """
        Sync Google Drive folder to local cache.
        
        Args:
            local_cache_dir: Local directory to cache files
            file_types: File types to sync
            
        Returns:
            List of synced file metadata
        """
        cache_path = Path(local_cache_dir)
        cache_path.mkdir(parents=True, exist_ok=True)
        
        files = self.list_files(file_types=file_types)
        synced_files = []
        
        for file_info in files:
            try:
                file_id = file_info['id']
                file_name = file_info['name']
                
                local_file = cache_path / f"{file_id}_{file_name}"
                
                # Download if not exists or modified
                if not local_file.exists():
                    content = self.download_file(file_id)
                    local_file.write_bytes(content)
                    logger.info(f"Synced: {file_name}")
                
                synced_files.append({
                    **file_info,
                    'local_path': str(local_file)
                })
                
            except Exception as e:
                logger.error(f"Failed to sync file {file_info['name']}: {e}")
        
        return synced_files


class FileParser:
    """
    Parse various file types to extract data for formulas.
    """
    
    @staticmethod
    def parse_pdf(file_path: str) -> Dict[str, Any]:
        """
        Extract text and tables from PDF.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted data with text and metadata
        """
        try:
            reader = PdfReader(file_path)
            
            text_content = []
            for page in reader.pages:
                text_content.append(page.extract_text())
            
            full_text = "\n\n".join(text_content)
            
            return {
                "type": "pdf",
                "pages": len(reader.pages),
                "text": full_text,
                "metadata": reader.metadata,
                "extracted_at": datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Failed to parse PDF {file_path}: {e}")
            raise
    
    @staticmethod
    def parse_docx(file_path: str) -> Dict[str, Any]:
        """
        Extract text and tables from DOCX.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text and tables
        """
        try:
            doc = DocxDocument(file_path)
            
            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                "type": "docx",
                "paragraphs": paragraphs,
                "tables": tables,
                "text": "\n".join(paragraphs),
                "extracted_at": datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Failed to parse DOCX {file_path}: {e}")
            raise
    
    @staticmethod
    def parse_excel(file_path: str) -> Dict[str, Any]:
        """
        Extract data from Excel file.
        
        Args:
            file_path: Path to Excel file
            
        Returns:
            Data from all sheets
        """
        try:
            # Read all sheets
            excel_data = pd.read_excel(file_path, sheet_name=None)
            
            sheets = {}
            for sheet_name, df in excel_data.items():
                sheets[sheet_name] = {
                    "columns": df.columns.tolist(),
                    "rows": len(df),
                    "data": df.to_dict('records'),
                    "summary": df.describe().to_dict() if len(df) > 0 else {}
                }
            
            return {
                "type": "excel",
                "sheets": list(sheets.keys()),
                "data": sheets,
                "extracted_at": datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Failed to parse Excel {file_path}: {e}")
            raise
    
    @staticmethod
    def parse_csv(file_path: str) -> Dict[str, Any]:
        """
        Extract data from CSV file.
        
        Args:
            file_path: Path to CSV file
            
        Returns:
            CSV data as dict
        """
        try:
            df = pd.read_csv(file_path)
            
            return {
                "type": "csv",
                "columns": df.columns.tolist(),
                "rows": len(df),
                "data": df.to_dict('records'),
                "summary": df.describe().to_dict() if len(df) > 0 else {},
                "extracted_at": datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Failed to parse CSV {file_path}: {e}")
            raise
    
    @staticmethod
    def parse_json(file_path: str) -> Dict[str, Any]:
        """
        Load JSON file.
        
        Args:
            file_path: Path to JSON file
            
        Returns:
            JSON data
        """
        try:
            with open(file_path) as f:
                data = json.load(f)
            
            return {
                "type": "json",
                "data": data,
                "extracted_at": datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Failed to parse JSON {file_path}: {e}")
            raise
    
    @classmethod
    def parse_file(cls, file_path: str) -> Dict[str, Any]:
        """
        Auto-detect file type and parse accordingly.
        
        Args:
            file_path: Path to file
            
        Returns:
            Parsed data
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        parsers = {
            '.pdf': cls.parse_pdf,
            '.docx': cls.parse_docx,
            '.doc': cls.parse_docx,
            '.xlsx': cls.parse_excel,
            '.xls': cls.parse_excel,
            '.csv': cls.parse_csv,
            '.json': cls.parse_json
        }
        
        parser = parsers.get(extension)
        if not parser:
            raise ValueError(f"Unsupported file type: {extension}")
        
        return parser(file_path)


class DataExtractor:
    """
    Extract formula inputs from parsed file data.
    """
    
    @staticmethod
    def extract_numerical_data(
        parsed_data: Dict[str, Any],
        patterns: Optional[Dict[str, str]] = None
    ) -> Dict[str, List[float]]:
        """
        Extract numerical values from parsed file data.
        
        Args:
            parsed_data: Output from FileParser
            patterns: Optional patterns to match specific values
            
        Returns:
            Dictionary of extracted numerical data
        """
        import re
        
        extracted = {}
        
        # Get text content
        text = parsed_data.get('text', '')
        
        # Extract all numbers with optional units
        number_pattern = r'(\w+)\s*[:=]\s*([\d.]+)\s*(\w+)?'
        matches = re.findall(number_pattern, text)
        
        for label, value, unit in matches:
            try:
                num_value = float(value)
                key = label.lower().strip()
                
                if key not in extracted:
                    extracted[key] = []
                
                extracted[key].append({
                    'value': num_value,
                    'unit': unit if unit else None
                })
            except ValueError:
                continue
        
        # Extract from tables if present
        if 'tables' in parsed_data:
            for table in parsed_data['tables']:
                if len(table) > 1:  # Has header
                    headers = table[0]
                    for row in table[1:]:
                        for i, value in enumerate(row):
                            try:
                                num_value = float(value)
                                key = headers[i].lower().strip()
                                
                                if key not in extracted:
                                    extracted[key] = []
                                
                                extracted[key].append({'value': num_value})
                            except (ValueError, IndexError):
                                continue
        
        # Extract from structured data (Excel, CSV)
        if 'data' in parsed_data and isinstance(parsed_data['data'], dict):
            for sheet_name, sheet_data in parsed_data['data'].items():
                if 'data' in sheet_data:
                    df_data = sheet_data['data']
                    for col in sheet_data.get('columns', []):
                        values = [row.get(col) for row in df_data]
                        numeric_values = []
                        for val in values:
                            try:
                                numeric_values.append({'value': float(val)})
                            except (ValueError, TypeError):
                                continue
                        
                        if numeric_values:
                            key = col.lower().strip()
                            extracted[key] = numeric_values
        
        return extracted
    
    @staticmethod
    def extract_context_hints(
        parsed_data: Dict[str, Any]
    ) -> Dict[str, str]:
        """
        Extract context hints from file content.
        
        Args:
            parsed_data: Output from FileParser
            
        Returns:
            Dictionary of context hints
        """
        import re
        
        text = parsed_data.get('text', '').lower()
        context = {}
        
        # Climate patterns
        climate_patterns = {
            'hot_arid': r'\b(hot|arid|desert|dry)\b',
            'hot_humid': r'\b(tropical|humid|monsoon)\b',
            'temperate': r'\b(temperate|moderate|mild)\b',
            'cold': r'\b(cold|arctic|freezing|winter)\b'
        }
        
        for climate_type, pattern in climate_patterns.items():
            if re.search(pattern, text):
                context['climate'] = climate_type
                break
        
        # Material patterns
        material_patterns = {
            'concrete': r'\b(concrete|cement)\b',
            'steel': r'\b(steel|iron)\b',
            'aluminum': r'\b(aluminum|aluminium)\b',
            'wood': r'\b(wood|timber)\b'
        }
        
        for material_type, pattern in material_patterns.items():
            if re.search(pattern, text):
                context['material'] = material_type
                break
        
        # Site conditions
        if re.search(r'\b(coastal|coast|beach|ocean)\b', text):
            context['site_condition'] = 'coastal'
        elif re.search(r'\b(mountain|elevation|altitude)\b', text):
            context['site_condition'] = 'mountain'
        
        # Project type
        if re.search(r'\b(building|construction|structure)\b', text):
            context['project_type'] = 'construction'
        elif re.search(r'\b(bridge)\b', text):
            context['project_type'] = 'bridge'
        
        return context


# Example usage
if __name__ == "__main__":
    # Google Drive sync
    connector = GoogleDriveConnector(
        credentials_path="/path/to/credentials.json",
        folder_id="your-folder-id"
    )
    
    files = connector.list_files(
        file_types=['application/pdf', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']
    )
    
    synced = connector.sync_folder(
        local_cache_dir="./drive_cache"
    )
    
    # File parsing
    for file_info in synced:
        local_path = file_info['local_path']
        
        # Parse file
        parsed = FileParser.parse_file(local_path)
        
        # Extract data
        numerical_data = DataExtractor.extract_numerical_data(parsed)
        context_hints = DataExtractor.extract_context_hints(parsed)
        
        print(f"File: {file_info['name']}")
        print(f"Extracted data: {numerical_data}")
        print(f"Context: {context_hints}")
